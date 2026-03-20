"""
JobPilot AI - Resume Parser
Extracts text and structure from PDF/DOCX resumes.
Then uses LLM to extract a structured CandidateProfile.
"""

import re
from pathlib import Path
from typing import Optional
from string import Template

from loguru import logger

from models.candidate import CandidateProfile, WorkExperience, Education, Project, Certification
from llm.client import LLMClient
from llm.prompts import RESUME_EXTRACTION_SYSTEM, RESUME_EXTRACTION_PROMPT


def extract_text_from_pdf(file_path: str) -> str:
    """Extract raw text from a PDF file using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("PyMuPDF not installed, trying pdfplumber...")
        return _extract_with_pdfplumber(file_path)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise


def _extract_with_pdfplumber(file_path: str) -> str:
    """Fallback PDF extractor using pdfplumber."""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text from a DOCX file."""
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_resume_text(file_path: str) -> str:
    """
    Auto-detect file type and extract text.
    Supports PDF and DOCX.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    suffix = path.suffix.lower()
    logger.info(f"Extracting text from {suffix} resume: {path.name}")

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif suffix == ".txt":
        return path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported resume format: {suffix}")


def clean_resume_text(text: str) -> str:
    """
    Clean extracted resume text:
    - Remove excessive whitespace
    - Remove non-printable characters
    - Normalize line breaks
    """
    # Normalize whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    # Remove non-printable chars
    text = re.sub(r'[^\x20-\x7E\n\t]', ' ', text)
    return text.strip()


class ResumeParser:
    """
    Main resume parser.
    Extracts text from file, then uses LLM to build a structured CandidateProfile.
    """

    def __init__(self, llm_client: Optional[LLMClient] = None):
        self.llm = llm_client or LLMClient()

    def parse(self, file_path: str) -> CandidateProfile:
        """
        Parse a resume file into a CandidateProfile.

        Args:
            file_path: Path to the resume (PDF or DOCX)

        Returns:
            CandidateProfile with all extracted information
        """
        logger.info(f"Parsing resume: {file_path}")

        # Step 1: Extract raw text
        raw_text = extract_resume_text(file_path)
        clean_text = clean_resume_text(raw_text)
        logger.debug(f"Extracted {len(clean_text)} characters from resume")

        # Step 2: Use LLM to parse into structured JSON
        prompt = Template(RESUME_EXTRACTION_PROMPT).substitute(
            resume_text=clean_text[:8000]  # Token limit safety
        )

        logger.info("Using LLM to extract structured profile...")
        parsed_data = self.llm.parse_json(
            user_prompt=prompt,
            system_prompt=RESUME_EXTRACTION_SYSTEM,
        )

        # Step 3: Build CandidateProfile
        profile = self._build_profile(parsed_data, raw_text, file_path)
        logger.success(
            f"Resume parsed: {profile.name} | "
            f"{profile.years_of_experience}y exp | "
            f"{len(profile.technical_skills)} skills"
        )
        return profile

    async def aparse(self, file_path: str) -> CandidateProfile:
        """Async version of parse."""
        raw_text = extract_resume_text(file_path)
        clean_text = clean_resume_text(raw_text)

        prompt = Template(RESUME_EXTRACTION_PROMPT).substitute(
            resume_text=clean_text[:8000]
        )

        parsed_data = await self.llm.aparse_json(
            user_prompt=prompt,
            system_prompt=RESUME_EXTRACTION_SYSTEM,
        )

        return self._build_profile(parsed_data, raw_text, file_path)

    def _build_profile(
        self,
        data: dict,
        raw_text: str,
        file_path: str,
    ) -> CandidateProfile:
        """Convert raw LLM JSON output into a CandidateProfile object."""
        # Build nested objects
        work_experiences = [
            WorkExperience(**exp)
            for exp in data.get("work_experience", [])
            if isinstance(exp, dict)
        ]
        educations = [
            Education(**edu)
            for edu in data.get("education", [])
            if isinstance(edu, dict)
        ]
        projects = [
            Project(**proj)
            for proj in data.get("projects", [])
            if isinstance(proj, dict)
        ]
        certifications = [
            Certification(**cert)
            for cert in data.get("certifications", [])
            if isinstance(cert, dict)
        ]

        return CandidateProfile(
            name=data.get("name", "Unknown"),
            email=data.get("email"),
            phone=data.get("phone"),
            location=data.get("location"),
            linkedin_url=data.get("linkedin_url"),
            github_url=data.get("github_url"),
            portfolio_url=data.get("portfolio_url"),
            summary=data.get("summary", ""),
            headline=data.get("headline", ""),
            years_of_experience=data.get("years_of_experience"),
            current_role=data.get("current_role"),
            current_company=data.get("current_company"),
            technical_skills=data.get("technical_skills", []),
            soft_skills=data.get("soft_skills", []),
            tools=data.get("tools", []),
            domains=data.get("domains", []),
            work_experience=work_experiences,
            education=educations,
            projects=projects,
            certifications=certifications,
            publications=data.get("publications", []),
            target_roles=data.get("target_roles", []),
            target_locations=data.get("target_locations", []),
            preferred_industries=data.get("preferred_industries", []),
            salary_expectation=data.get("salary_expectation"),
            notice_period=data.get("notice_period"),
            remote_preference=data.get("remote_preference"),
            raw_resume_text=raw_text,
            resume_file_path=file_path,
        )
